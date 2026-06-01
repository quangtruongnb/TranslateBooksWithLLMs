"""Small, deterministic input fixtures for the four supported formats.

Each builder writes a file into ``dest_dir`` and returns its path. Content is
sized to produce several chunks at ``max_tokens_per_chunk=60`` so multi-chunk
progress is exercised.
"""

import importlib.util
from pathlib import Path

_REPO_ROOT = Path(__file__).resolve().parents[2]


def _para(n: int) -> str:
    return (
        f"This is paragraph number {n}. "
        "It contains several sentences so that the token chunker has enough "
        "material to split the document into more than one chunk. "
        "The quick brown fox jumps over the lazy dog while the patient "
        "translator carefully preserves the meaning of every clause."
    )


def build_txt(dest_dir: Path) -> Path:
    path = dest_dir / "sample.txt"
    body = "\n\n".join(_para(i) for i in range(1, 7))
    path.write_text(body, encoding="utf-8")
    return path


def build_srt(dest_dir: Path) -> Path:
    # 50 subtitles at the default 20 lines/block yields 3 blocks, so the
    # multi-chunk progression (and the "blocks, not subtitles" counting) is
    # exercised rather than collapsing to a single block.
    path = dest_dir / "sample.srt"
    blocks = []
    for i in range(1, 51):
        total_s = i * 3
        start = f"00:{total_s // 60:02d}:{total_s % 60:02d},000"
        end_s = total_s + 2
        end = f"00:{end_s // 60:02d}:{end_s % 60:02d},500"
        blocks.append(
            f"{i}\n{start} --> {end}\nSubtitle line {i}: the fox runs swiftly.\n"
        )
    path.write_text("\n".join(blocks) + "\n", encoding="utf-8")
    return path


def build_docx(dest_dir: Path) -> Path:
    import docx  # python-docx, available in the dev environment

    path = dest_dir / "sample.docx"
    document = docx.Document()
    document.add_heading("Characterization Sample", level=1)
    # Enough paragraphs to span several HTML chunks so multi-chunk DOCX
    # progress is exercised, not just a single 0->done callback.
    for i in range(1, 21):
        document.add_paragraph(_para(i))
    document.save(str(path))
    return path


def build_epub(dest_dir: Path) -> Path:
    """Build the canonical Translator's Sampler EPUB (~10 chunks)."""
    spec = importlib.util.spec_from_file_location(
        "scripts.build_test_epub", _REPO_ROOT / "scripts" / "build_test_epub.py"
    )
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    path = dest_dir / "sample.epub"
    module.build_epub(path)
    return path
